#!/usr/bin/env python3
"""Markdown → DOCX converter for the resume composer pipeline.

Converts a markdown resume (produced by `resume_compose.py`) into a styled
Word document matching the user's preferred format. Originally lived in a
separate resume-converter repo; vendored into Finder to eliminate the
cross-repo path dependency and keep the whole resume pipeline self-contained.

Markdown format expected (matches what `resume_compose.py` emits):

    # Name
    Contact info line (centered)

    Summary / archetype intro paragraph

    ## SECTION HEADER

    ### Company Name | Location
    #### Job Title | Dates
    - Bullet point
    - Bullet point

    ## TECHNICAL SKILLS

    **Category:** Skill 1, Skill 2, Skill 3

    ---
    Footer text (centered, optional)

Output formatting:
- Name: 16pt Bold, ALL CAPS, centered, with horizontal rule under it
- Contact: 10pt centered
- Summary: 10.5pt plain paragraph
- Section headers: 12pt Bold Dark Blue (#1F4E78), ALL CAPS
- Company: 11pt Bold (location right-tab-aligned to margin)
- Title: 10.5pt (dates right-tab-aligned)
- Bullets: `•` + tab + content, hanging indent for wrap alignment
- Skills: bold "Category:" label + plain content
- Footer: 10pt centered (only emitted if `---` separator present)
- Margins: 0.5" all sides
- Body default: 10.5pt (tight for 2-page fit)

Requires python-docx (see requirements.txt).
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_document_defaults(doc: Document) -> None:
    """Set up document margins, defaults, and base font size for 2-page fit."""
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    doc.styles['Normal'].font.size = Pt(10.5)


def add_name(doc: Document, text: str) -> None:
    """Add the name at the top (16pt Bold, Centered) with a horizontal rule under it."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(2)

    # Bottom border = horizontal rule directly under name (above contact line)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')      # 0.75pt line weight (1/8 pt units)
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_contact(doc: Document, text: str) -> None:
    """Add contact info line (10pt, Centered). No border (lives on name above)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_summary(doc: Document, text: str) -> None:
    """Add summary paragraph."""
    p = doc.add_paragraph()
    p.add_run(text)
    p.paragraph_format.space_after = Pt(6)


def add_section_header(doc: Document, text: str) -> None:
    """Add section header (12pt Bold, Dark Blue #1F4E78)."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)


def add_company_line(doc: Document, company: str, location: str) -> None:
    """Add company/organization line with location right-aligned via tabs."""
    p = doc.add_paragraph()

    # Right-aligned tab at the right margin.
    # Page width 8.5" - margins (0.5" each side) = 7.5" usable width
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)

    run = p.add_run(company)
    run.bold = True
    run.font.size = Pt(11)

    run2 = p.add_run(f"\t{location}")
    run2.font.size = Pt(11)

    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)


def add_title_line(doc: Document, title: str, dates: str) -> None:
    """Add job title line with dates right-aligned via tabs."""
    p = doc.add_paragraph()
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run(f"{title}\t{dates}")
    p.paragraph_format.space_after = Pt(2)


def add_bullet(doc: Document, text: str) -> None:
    """Add a bullet point with explicit • glyph + tab-stop hanging indent.

    python-docx's `style='List Paragraph'` is indent-only. It does NOT
    auto-render the bullet glyph unless the numbering.xml part is also
    configured (which python-docx doesn't ship by default). Using a literal
    bullet glyph + a tab stop at the body indent position guarantees wrapped
    continuation lines align exactly with the first-line content (not with
    the bullet glyph itself).
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(0.2))
    p.add_run('•\t')
    add_formatted_text(p, text)
    p.paragraph_format.space_after = Pt(1)


def add_skill_line(doc: Document, category: str, content: str) -> None:
    """Add a skills line with bold category."""
    p = doc.add_paragraph()
    run = p.add_run(f"{category}: ")
    run.bold = True
    p.add_run(content)
    p.paragraph_format.space_after = Pt(1)


def add_plain_paragraph(doc: Document, text: str) -> None:
    """Add a plain paragraph with optional inline formatting."""
    p = doc.add_paragraph()
    add_formatted_text(p, text)
    p.paragraph_format.space_after = Pt(2)


def add_footer(doc: Document, text: str) -> None:
    """Add centered footer text (10pt)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_before = Pt(6)


def add_formatted_text(paragraph, text: str) -> None:
    """Parse and add text with **bold** and *italic* formatting."""
    pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))'
    for match in re.finditer(pattern, text):
        if match.group(2):       # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):     # **bold**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):     # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):     # plain
            paragraph.add_run(match.group(5))


def convert_markdown_to_word(md_path: str, output_path: str) -> None:
    """Convert markdown resume to Word document."""
    lines = Path(md_path).read_text(encoding='utf-8').splitlines()
    doc = Document()
    set_document_defaults(doc)

    state = 'start'  # start, after_name, after_contact, body
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue

        # H1: Name (only at document start)
        if line.startswith('# ') and state == 'start':
            add_name(doc, line[2:].strip())
            state = 'after_name'
            i += 1
            continue

        # Contact line (first non-empty line after name)
        if state == 'after_name' and not line.startswith('#'):
            add_contact(doc, line)
            state = 'after_contact'
            i += 1
            continue

        # Summary (first paragraph after contact, before any ##)
        if state == 'after_contact' and not line.startswith('#'):
            summary_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#'):
                summary_lines.append(lines[i].strip())
                i += 1
            add_summary(doc, ' '.join(summary_lines))
            state = 'body'
            continue

        if line.startswith('## '):
            add_section_header(doc, line[3:].strip())
            state = 'body'
            i += 1
            continue

        if line.startswith('### '):
            content = line[4:].strip()
            if '|' in content:
                parts = content.split('|', 1)
                add_company_line(doc, parts[0].strip(), parts[1].strip())
            else:
                add_company_line(doc, content, '')
            i += 1
            continue

        if line.startswith('#### '):
            content = line[5:].strip()
            if '|' in content:
                parts = content.split('|', 1)
                add_title_line(doc, parts[0].strip(), parts[1].strip())
            else:
                add_title_line(doc, content, '')
            i += 1
            continue

        if line.startswith('- ') or line.startswith('* '):
            add_bullet(doc, line[2:].strip())
            i += 1
            continue

        # Horizontal rule → next non-empty line is the footer
        if line.startswith('---'):
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            if i < len(lines):
                add_footer(doc, lines[i].strip())
                i += 1
            continue

        # Skill line: **Category:** content
        skill_match = re.match(r'\*\*(.+?):\*\*\s*(.+)', line)
        if skill_match:
            add_skill_line(doc, skill_match.group(1), skill_match.group(2))
            i += 1
            continue

        # Plain paragraph fallback
        if state == 'body':
            add_plain_paragraph(doc, line)
        i += 1

    doc.save(output_path)
    print(f"Created: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Convert a markdown resume to a Word document.')
    parser.add_argument('input', help='Path to markdown file')
    parser.add_argument('-o', '--output', help='Output path (default: same name with .docx extension)')
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERROR: File not found: {input_path}", file=sys.stderr)
        return 2

    output_path = args.output or str(input_path.with_suffix('.docx'))
    convert_markdown_to_word(str(input_path), output_path)
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
