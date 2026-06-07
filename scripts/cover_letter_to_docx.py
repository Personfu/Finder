#!/usr/bin/env python3
"""Cover letter markdown -> styled DOCX.

Companion to scripts/md_to_docx.py (which handles the resume). The cover
letter has a different shape (no section headers, just date + paragraphs)
so it gets its own renderer. Same letterhead style as the resume.

Markdown format expected (matches what /finder:apply writes):

    <three paragraphs of body content, separated by blank lines>

The name + contact + date letterhead is rendered programmatically; the
markdown only carries the body.

Output formatting matches md_to_docx.py's resume header:
- Name: 16pt Bold, ALL CAPS, centered, horizontal rule under it
- Contact: 10pt centered
- Date: 10.5pt left-aligned
- Body paragraphs: 10.5pt with 10pt space-after
- Margins: 0.5" all sides

Requires python-docx (same dep as md_to_docx.py).
"""

from __future__ import annotations

import argparse
import datetime as _dt
import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _load_canon() -> dict:
    """Read name + contact from data/canon.json (single source of truth)."""
    repo_root = Path(__file__).resolve().parent.parent
    canon_path = repo_root / "data" / "canon.json"
    if not canon_path.is_file():
        print(
            f"WARNING: {canon_path} not found; cover-letter letterhead will use "
            f"placeholder name/contact. Populate data/canon.json (see "
            f"data/canon.example.json) for a real letterhead.",
            file=sys.stderr,
        )
        return {}
    return json.loads(canon_path.read_text(encoding="utf-8"))


def _name_and_contact(canon: dict) -> tuple[str, str]:
    """Build the (name, contact-line) pair from the canon `personal` block.

    Mirrors the schema used by resume_compose.py (name / location / phone /
    email_security|email_personal) so the cover-letter letterhead matches the
    resume exactly. Fallbacks are generic placeholders, never real PII.
    """
    personal = canon.get("personal", {}) if isinstance(canon, dict) else {}
    name = personal.get("name") or "Firstname Lastname"
    location = personal.get("location") or "City, ST ZIP"
    phone = personal.get("phone") or "555-555-5555"
    email = (
        personal.get("email_security")
        or personal.get("email_personal")
        or "firstname@example.com"
    )
    contact = f"{location} | {phone} | {email}"
    return name, contact


def _set_defaults(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    doc.styles["Normal"].font.size = Pt(10.5)


def _add_name(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(16)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_contact(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(14)


def _add_date(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    p.paragraph_format.space_after = Pt(14)


def _add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.add_run(text)
    p.paragraph_format.space_after = Pt(10)


def convert(cover_md_path: Path, out_path: Path, today: _dt.date | None = None) -> None:
    body_md = cover_md_path.read_text(encoding="utf-8").strip()
    if not body_md or body_md.startswith("<!--"):
        raise ValueError(
            f"refusing to convert {cover_md_path}: file is empty or still a placeholder. "
            f"Run /finder:apply first to draft the cover letter."
        )

    paragraphs = [chunk.strip() for chunk in body_md.split("\n\n") if chunk.strip()]
    if not paragraphs:
        raise ValueError(f"no body paragraphs found in {cover_md_path}")

    canon = _load_canon()
    name, contact = _name_and_contact(canon)
    today = today or _dt.date.today()
    date_str = today.strftime("%B %d, %Y")

    doc = Document()
    _set_defaults(doc)
    _add_name(doc, name)
    _add_contact(doc, contact)
    _add_date(doc, date_str)
    for para in paragraphs:
        # Collapse internal line breaks within a paragraph; markdown
        # convention is single blank line between paragraphs.
        single_line = " ".join(line.strip() for line in para.splitlines() if line.strip())
        _add_body_paragraph(doc, single_line)

    doc.save(str(out_path))


def main() -> int:
    p = argparse.ArgumentParser(description="Convert a cover letter markdown to a styled DOCX.")
    p.add_argument("input", help="Path to cover_letter.md (or a packet directory containing it)")
    p.add_argument("-o", "--output", help="Output DOCX path (default: <input>.docx)")
    p.add_argument(
        "--today",
        default=None,
        help="Override today's date for the letterhead (YYYY-MM-DD). Useful for reproducible runs.",
    )
    args = p.parse_args()

    input_path = Path(args.input)
    if input_path.is_dir():
        input_path = input_path / "cover_letter.md"
    if not input_path.exists():
        print(f"ERROR: cover letter markdown not found: {input_path}", file=sys.stderr)
        return 2

    if args.output:
        out_path = Path(args.output)
    else:
        out_path = input_path.with_suffix(".docx")

    today = None
    if args.today:
        try:
            today = _dt.date.fromisoformat(args.today)
        except ValueError as e:
            print(f"ERROR: bad --today value: {e}", file=sys.stderr)
            return 2

    try:
        convert(input_path, out_path, today=today)
    except ValueError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 2
    print(f"Created: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
